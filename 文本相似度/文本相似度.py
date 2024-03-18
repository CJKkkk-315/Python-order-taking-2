from gensim import corpora, models, similarities
import jieba
import os
import docx
# 输入要比对的两个目标文件
file1_name = '1.docx'
file2_name = '2.docx'
texts = []
# 设置要读取的文件后缀为docx
ext = 'docx'
files = os.listdir(os.getcwd())
# 读取当前目录下所有docx文件
files = [i for i in files if i.split('.')[-1] == ext]
# 遍历每个docx文件
for file in files:
    # 读取其中的所有文本以及表格中的文本
    file = docx.Document(file)
    content = [i.text for i in file.paragraphs]
    for i in file.tables:
        for j in i.rows:
            for k in j.cells:
                content.append(k.text)
    # 构建成字符串，加入到texts列表中
    content = ''.join(content)
    texts.append(content)
# 提取中要比对的第一个文件
keyword = texts[files.index(file1_name)]
# 将字符串列表分词生成分词列表
texts = [jieba.lcut(text) for text in texts]
# 基于分词列表建立词语词典，并提取词典特征数
dictionary = corpora.Dictionary(texts)
feature_cnt = len(dictionary.token2id)
# 基于词典，将分词列表集转换成稀疏向量语料库
corpus = [dictionary.doc2bow(text) for text in texts]
# 使用gensim库中的TF-IDF模型生成语料库
tfidf = models.TfidfModel(corpus)
# 对tfidf生成的语料库建立索引
index = similarities.SparseMatrixSimilarity(tfidf[corpus], num_features=feature_cnt)
# 对要比对的文件1也做同样的处理
kw_vector = dictionary.doc2bow(jieba.lcut(keyword))
# 相似度计算
sim = index[tfidf[kw_vector]][files.index(file2_name)]
print(round(sim,2))